from appium import webdriver
import loginmodule_social
import timemodule
import connectmodule_social
from docx import Document
from wordtest import doc_init
import os
import selenium
#import time
import appium
import sys
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

current_time = timemodule.current_time()
print_time = timemodule.print_time()

doc = Document()
report_header = doc.add_paragraph('Login Test Report: We think the time is ')
report_header.add_run(current_time)
report_header.alignment = WD_ALIGN_PARAGRAPH.CENTER

print(print_time)

table = doc.add_table(rows=3, cols=1)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

print("Login Test, don't crash and burn pls")
print("Mod imported")
cell00 = table.cell(0, 0)
cell00.text = "Mod Imported"

desired_cap ={
    "platformName": "Android",
    "deviceName": "Galaxy S7 edge",
    "platformVersion": "8",
    "app": "D:\\Documents\\Oscar\\DNANudge\\app-social-release.apk",
    "appPackage": "com.dnanudge.android.social",
    "appActivity": "com.dnanudge.android.features.splashScreen.SplashActivity",
    "fullReset": "False",
    "noReset": "True"
}

print("Cap loaded")
cell10 = table.cell(1, 0)
cell10.text = ('Cap loaded')
driver = webdriver.Remote("http://localhost:4723/wd/hub", desired_cap)
print("Driver loaded")
cell20 = table.cell(2, 0)
cell20.text = ('Driver loaded')

#####SETTINGS#####

login = False
login_pass = True
login_skip = False
connect = True
connect_pass = True
connect_setup = True

#^^^^SETTINGS^^^^#

if login == True and login_pass == True and login_skip == True:
    loginmodule_social.login_pass_44(driver)
    loginmodule_social.login_skip(driver)
elif login == True and login_pass == True and login_skip == False:
    loginmodule_social.login_pass_44(driver)
elif login == True and login_pass == False:
    loginmodule_social.login_fail(driver)
else:
    pass

if connect == True and connect_pass == True and connect_setup == True:
    connectmodule_social.find_lifestyle(driver)
    connectmodule_social.find_dnaband(driver)
    connectmodule_social.find_dnaband_pass(driver)
    connectmodule_social.setup_dnaband(driver)
    connectmodule_social.tutorial_dnaband(driver)
elif connect == True and connect_pass == True and connect_setup == False:
    connectmodule_social.find_lifestyle(driver)
    connectmodule_social.find_dnaband(driver)
    connectmodule_social.find_dnaband_pass(driver)
    connectmodule_social.tutorial_dnaband(driver)
elif connect == True and connect_pass == False:
    connectmodule_social.find_lifestyle(driver)
    connectmodule_social.find_dnaband(driver)
    connectmodule_social.find_dnaband_fail(driver)
    connectmodule_social.find_dnaband_fail_loop(driver, 1)
else:
    connectmodule_social.find_lifestyle(driver)

print("It worked! Praise the gods!")

success_response = doc.add_paragraph('It worked! Praise the')
success_response.add_run(" gods!").bold=True
#doc.save(print_time + '_demo.docx')
#os.system(print_time + '_demo.docx')